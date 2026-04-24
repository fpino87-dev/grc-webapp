"""
Markdown → .docx converter for GRC procedure documents.
Handles: headings H1-H4, bullet lists, numbered lists, inline bold/italic/code.
"""

import io
import re

_INLINE_RE = re.compile(
    r"(\*\*\*[^*]+?\*\*\*|\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)"
)


def _parse_inline(paragraph, text: str) -> None:
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("***") and part.endswith("***"):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Courier New"
        else:
            run.text = part


def markdown_to_docx(md_text: str, title: str = "") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in md_text.splitlines():
        line = line.rstrip()

        if not line.strip():
            continue

        stripped = line.lstrip()
        indent = len(line) - len(stripped)

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            _parse_inline(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, re.sub(r"^\d+\.\s+", "", stripped).strip())
        elif re.match(r"^-{3,}$", stripped):
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _parse_inline(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
